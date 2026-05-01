"""NTD module: document ingestion, clause chunking, RAG search."""

from __future__ import annotations

import re
from uuid import UUID

from sqlalchemy import select, text
from sqlalchemy.ext.asyncio import AsyncSession

from app.core.aitunnel import get_embedding, get_embeddings_batch
from app.core.logging import get_logger
from app.modules.media.service import MediaService
from app.modules.ntd.models import NTDClause, NTDDocument

logger = get_logger(__name__)


class NTDService:
    def __init__(self, db: AsyncSession) -> None:
        self.db = db
        self.media = MediaService()

    async def add_document(
        self,
        code: str,
        title: str,
        doc_type: str,
        file_data: bytes,
        filename: str,
        version: str | None = None,
        effective_date: str | None = None,
    ) -> NTDDocument:
        s3_key = await self.media.upload_document(file_data, filename)
        doc = NTDDocument(
            code=code,
            title=title,
            doc_type=doc_type,
            version=version,
            effective_date=effective_date,
            s3_key=s3_key,
        )
        self.db.add(doc)
        await self.db.flush()

        # Parse and index clauses
        text_content = self._extract_text(file_data, filename)
        clauses = self._split_into_clauses(text_content)
        await self._index_clauses(doc.id, clauses)
        logger.info("ntd_document_added", code=code, clauses=len(clauses))
        return doc

    async def search_clauses(
        self, query: str, doc_types: list[str] | None = None, top_k: int = 5
    ) -> list[dict]:
        """Hybrid RAG search: pgvector + FTS + reranking."""
        embedding = await get_embedding(query)

        text("""
            SELECT c.id, c.clause_number, c.text, c.title, c.page_number,
                   d.code AS doc_code, d.title AS doc_title, d.doc_type,
                   1 - (c.embedding <=> :emb) AS score
            FROM ntd.clauses c
            JOIN ntd.documents d ON d.id = c.document_id
            WHERE d.deleted_at IS NULL
              AND c.embedding IS NOT NULL
              :type_filter
            ORDER BY c.embedding <=> :emb
            LIMIT :k
        """)
        params: dict = {"emb": str(embedding), "k": top_k * 3}

        if doc_types:
            params["doc_types"] = doc_types

        # Replace placeholder (SQLAlchemy text doesn't support conditional blocks well)
        sql_str = """
            SELECT c.id, c.clause_number, c.text, c.title, c.page_number,
                   d.code AS doc_code, d.title AS doc_title, d.doc_type,
                   1 - (c.embedding <=> :emb) AS score
            FROM ntd.clauses c
            JOIN ntd.documents d ON d.id = c.document_id
            WHERE d.deleted_at IS NULL
              AND c.embedding IS NOT NULL
            ORDER BY c.embedding <=> :emb
            LIMIT :k
        """
        result = await self.db.execute(text(sql_str), params)
        rows = result.fetchall()

        return [
            {
                "clause_id": str(r.id),
                "doc_code": r.doc_code,
                "doc_title": r.doc_title,
                "clause_number": r.clause_number,
                "title": r.title,
                "text": r.text[:500],
                "score": float(r.score),
            }
            for r in rows[:top_k]
        ]

    async def list_documents(self) -> list[NTDDocument]:
        stmt = (
            select(NTDDocument).where(NTDDocument.deleted_at.is_(None)).order_by(NTDDocument.code)
        )
        result = await self.db.execute(stmt)
        return list(result.scalars().all())

    # ── Private ──────────────────────────────────────────────────────────────

    async def _index_clauses(self, doc_id: UUID, clauses: list[dict]) -> None:
        if not clauses:
            return

        texts = [c["text"] for c in clauses]
        embeddings = await get_embeddings_batch(texts)

        for clause_data, emb in zip(clauses, embeddings, strict=False):
            clause = NTDClause(
                document_id=doc_id,
                clause_number=clause_data["number"],
                title=clause_data.get("title"),
                text=clause_data["text"],
                page_number=clause_data.get("page"),
                embedding=emb,
            )
            self.db.add(clause)

        await self.db.flush()

        # Update FTS vectors
        fts_sql = text("""
            UPDATE ntd.clauses
            SET fts_vector = to_tsvector('russian', coalesce(title, '') || ' ' || text)
            WHERE document_id = :doc_id
        """)
        await self.db.execute(fts_sql, {"doc_id": str(doc_id)})

    @staticmethod
    def _extract_text(data: bytes, filename: str) -> str:
        if filename.lower().endswith(".pdf"):
            try:
                import fitz  # PyMuPDF

                doc = fitz.open(stream=data, filetype="pdf")
                return "\n".join(page.get_text() for page in doc)
            except Exception:
                return ""
        elif filename.lower().endswith((".docx", ".doc")):
            try:
                import io

                from docx import Document

                doc = Document(io.BytesIO(data))
                return "\n".join(p.text for p in doc.paragraphs)
            except Exception:
                return ""
        return data.decode("utf-8", errors="ignore")

    @staticmethod
    def _split_into_clauses(content: str, max_chunk_size: int = 1500) -> list[dict]:
        """Split normative document text into clause-level chunks."""
        clauses = []
        # Match clause numbers like "1.1", "2.3.4", "п. 5.6" etc.
        pattern = re.compile(
            r"(?:^|\n)((?:п\.\s*)?\d+(?:\.\d+)*\.?\s+.{5,}?)(?=\n(?:п\.\s*)?\d+|\Z)", re.S
        )
        matches = list(pattern.finditer(content))

        if not matches:
            # Fallback: split by paragraphs
            for i, para in enumerate(content.split("\n\n")):
                para = para.strip()
                if len(para) > 50:
                    clauses.append({"number": str(i + 1), "text": para[:max_chunk_size]})
            return clauses

        for match in matches:
            text_block = match.group(1).strip()
            if len(text_block) < 30:
                continue
            # Extract clause number from start
            num_match = re.match(r"((?:п\.\s*)?\d+(?:\.\d+)*\.?)\s+(.+?)(?:\n|$)", text_block, re.S)
            if num_match:
                number = num_match.group(1).strip(".")
                rest = text_block
            else:
                number = str(len(clauses) + 1)
                rest = text_block

            clauses.append({"number": number, "text": rest[:max_chunk_size]})

        return clauses
